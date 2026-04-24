import os
import time
import logging
import requests
import uuid
import json
import zipfile
import pandas as pd
from io import BytesIO
from docx import Document
from flask import Flask, request, render_template, send_file, flash, redirect, url_for, session, jsonify, get_flashed_messages
from werkzeug.utils import secure_filename
from datetime import timedelta
from pathlib import Path

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'bmp', 'tif', 'tiff', 'webp', 'zip'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 64 * 1024 * 1024  # 64 MB limit
app.secret_key = os.environ.get('SECRET_KEY', 'default-secret-key')
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=1)

PADDLEOCR_API_URL = os.environ.get('PADDLEOCR_API_URL', 'http://paddleocr-api:8000')

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

def allowed_file(filename):
    if not filename:
        return False
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    tasks = session.get('tasks', [])
    return render_template('index.html', has_tasks=len(tasks) > 0)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('Файл не найден', 'error')
        return redirect(url_for('index'))

    uploaded_files = request.files.getlist('file')
    if not uploaded_files or (len(uploaded_files) == 1 and uploaded_files[0].filename == ''):
        flash('Файлы не выбраны', 'error')
        return redirect(url_for('index'))

    if 'tasks' not in session:
        session['tasks'] = []
    
    processed_any = False
    detect_seal = request.form.get('detect_seal') == 'on'
    
    for file in uploaded_files:
        if not file.filename or not allowed_file(file.filename):
            continue

        filename = secure_filename(file.filename)
        
        # Check if it's a ZIP file
        if filename.lower().endswith('.zip'):
            try:
                zip_data = BytesIO(file.read())
                with zipfile.ZipFile(zip_data) as z:
                    for zinfo in z.infolist():
                        if zinfo.is_dir():
                            continue
                        
                        z_filename = os.path.basename(zinfo.filename)
                        if not z_filename or z_filename.startswith('.') or not allowed_file(z_filename) or z_filename.lower().endswith('.zip'):
                            continue
                        
                        with z.open(zinfo) as zf:
                            file_content = zf.read()
                            
                            # Send to OCR API
                            files = {'file': (z_filename, file_content)}
                            data = {'detect_seal': detect_seal}
                            try:
                                response = requests.post(f"{PADDLEOCR_API_URL}/ocr", files=files, data=data)
                                response.raise_for_status()
                                job_data = response.json()
                                
                                task_id = job_data['job_id']
                                session['tasks'].append({'task_id': task_id, 'filename': z_filename})
                                processed_any = True
                            except Exception as e:
                                logger.error(f"Error sending file {z_filename} from ZIP to OCR API: {e}")
            except Exception as e:
                logger.error(f"Error processing ZIP file: {e}")
                flash(f"Ошибка при обработке ZIP архива {filename}: {e}", 'error')
        else:
            # Send file to PaddleOCR API
            try:
                files = {'file': (filename, file.read(), file.content_type)}
                data = {'detect_seal': detect_seal}
                response = requests.post(f"{PADDLEOCR_API_URL}/ocr", files=files, data=data)
                response.raise_for_status()
                job_data = response.json()
                
                task_id = job_data['job_id']
                session['tasks'].append({'task_id': task_id, 'filename': filename})
                processed_any = True
            except Exception as e:
                logger.error(f"Error sending file to PaddleOCR API: {e}")
                flash(f"Ошибка подключения к OCR бэкенду для {filename}: {e}", 'error')

    session.modified = True
    if not processed_any:
        if not get_flashed_messages():
            flash('Не удалось обработать ни один файл', 'error')
        return redirect(url_for('index'))
    
    return redirect(url_for('status_page'))

@app.route('/status')
def status_page():
    tasks = session.get('tasks', [])
    if not tasks:
        return redirect(url_for('index'))
    return render_template('status.html', tasks=tasks)

@app.route('/status/<task_id>')
def status(task_id):
    # This route is now mostly for backward compatibility or direct access
    # but we can make it redirect to the new status page or handle it there
    return render_template('status.html', tasks=[{'task_id': task_id, 'filename': 'Документ'}], single_task=task_id)

@app.route('/api/task_status/<task_id>')
def task_status(task_id):
    try:
        response = requests.get(f"{PADDLEOCR_API_URL}/ocr/{task_id}")
        response.raise_for_status()
        job_data = response.json()
        
        # Map PaddleOCR API status to what frontend expects
        # Frontend expects: status, progress, step, redirect, error
        status_map = {
            'queued': 'processing',
            'processing': 'processing',
            'completed': 'completed',
            'failed': 'failed',
            'cancelled': 'failed'
        }
        
        status = status_map.get(job_data['status'], 'processing')
        
        progress = 0
        if job_data['total_pages'] > 0:
            progress = int((job_data['processed_pages'] / job_data['total_pages']) * 100)
        elif job_data['status'] == 'completed':
            progress = 100
            
        result = {
            "status": status,
            "progress": progress,
            "step": "ocr" if status == "processing" else "",
            "total_pages": job_data['total_pages'],
            "processed_pages": job_data['processed_pages']
        }
        
        if status == 'completed':
            result["redirect"] = url_for('success', task_id=task_id)
            # Calculate processing time
            if 'created_at' in job_data and 'updated_at' in job_data:
                processing_time = job_data['updated_at'] - job_data['created_at']
                session[f'processing_time_{task_id}'] = round(processing_time, 2)
        
        if status == 'failed':
            result["error"] = job_data.get('error', 'Неизвестная ошибка')
            result["redirect"] = url_for('index')
            
        return jsonify(result)
    except Exception as e:
        logger.error(f"Error checking status: {e}")
        return jsonify({"status": "failed", "error": str(e), "redirect": url_for('index')})

@app.route('/success/<task_id>')
def success(task_id):
    tasks = session.get('tasks', [])
    filename = next((t['filename'] for t in tasks if t['task_id'] == task_id), 'документ')
    processing_time = session.get(f'processing_time_{task_id}', 'Н/Д')
    detect_seal_enabled = False
    
    # Auto-save results to output folder
    try:
        response = requests.get(f"{PADDLEOCR_API_URL}/ocr/{task_id}")
        if response.status_code == 200:
            job_data = response.json()
            detect_seal_enabled = bool(job_data.get('detect_seal', 0))

        response = requests.get(f"{PADDLEOCR_API_URL}/ocr/{task_id}/result")
        if response.status_code == 200:
            result_data = response.json()
            
            # Save as JSON
            json_filename = f"{task_id}.json"
            with open(os.path.join(app.config['OUTPUT_FOLDER'], json_filename), "w", encoding="utf-8") as f:
                json.dump(result_data, f, ensure_ascii=False, indent=2)
            
            # Save as Markdown
            md_filename = f"{task_id}.md"
            pages = result_data.get('pages', [])
            full_markdown = ""
            for page in pages:
                full_markdown += f"# Страница {page['page_num']}\n\n{page['markdown']}\n\n---\n\n"
            
            with open(os.path.join(app.config['OUTPUT_FOLDER'], md_filename), "w", encoding="utf-8") as f:
                f.write(full_markdown)
                
            logger.info(f"Results for task {task_id} saved to {app.config['OUTPUT_FOLDER']}")
    except Exception as e:
        logger.error(f"Failed to auto-save results for task {task_id}: {e}")

    return render_template('success.html', task_id=task_id, filename=filename, processing_time=processing_time, detect_seal_enabled=detect_seal_enabled)

@app.route('/api/image/<task_id>/<int:page_num>')
def get_image(task_id, page_num):
    try:
        response = requests.get(f"{PADDLEOCR_API_URL}/ocr/{task_id}/image/{page_num}", stream=True)
        response.raise_for_status()
        return send_file(response.raw, mimetype='image/png')
    except Exception as e:
        logger.error(f"Error fetching image: {e}")
        return "Изображение не найдено", 404

@app.route('/api/seals/<task_id>')
def list_seals(task_id):
    try:
        response = requests.get(f"{PADDLEOCR_API_URL}/ocr/{task_id}/seals")
        response.raise_for_status()
        return jsonify(response.json())
    except Exception as e:
        logger.error(f"Error listing seals: {e}")
        return jsonify({"seals": []})

@app.route('/api/seal/<task_id>/<filename>')
def get_seal(task_id, filename):
    try:
        response = requests.get(f"{PADDLEOCR_API_URL}/ocr/{task_id}/seals/{filename}", stream=True)
        response.raise_for_status()
        return send_file(response.raw, mimetype='image/png')
    except Exception as e:
        logger.error(f"Error fetching seal: {e}")
        return "Печать не найдена", 404

@app.route('/download_result/<task_id>')
def download_result(task_id):
    fmt = request.args.get('format', 'md').lower()
    try:
        response = requests.get(f"{PADDLEOCR_API_URL}/ocr/{task_id}/result")
        response.raise_for_status()
        result_data = response.json()
        
        pages = result_data.get('pages', [])
        
        if fmt == 'docx':
            doc = Document()
            for page in pages:
                doc.add_heading(f"Страница {page['page_num']}", level=1)
                # Split markdown by newlines and add each non-empty line as a separate paragraph
                markdown_text = page.get('markdown', '')
                for line in markdown_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                doc.add_page_break()
            
            mem = BytesIO()
            doc.save(mem)
            mem.seek(0)
            return send_file(
                mem,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"ocr_result_{task_id}.docx"
            )
        
        elif fmt == 'xlsx':
            rows = []
            for page in pages:
                if 'result_json' in page and page['result_json']:
                    # result_json is a list of pages
                    for p_data in page['result_json']:
                        for block in p_data.get('blocks', []):
                            rows.append({
                                'Страница': page['page_num'],
                                'Текст': block.get('text', ''),
                                'Уверенность': block.get('confidence', 0)
                            })
                else:
                    # Fallback to markdown if result_json is missing
                    rows.append({
                        'Страница': page['page_num'],
                        'Текст': page['markdown'],
                        'Уверенность': 1.0
                    })
            
            df = pd.DataFrame(rows)
            mem = BytesIO()
            with pd.ExcelWriter(mem, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='OCR Result')
            mem.seek(0)
            return send_file(
                mem,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f"ocr_result_{task_id}.xlsx"
            )
            
        else: # Default to MD
            # Extract full markdown
            full_markdown = ""
            for page in pages:
                full_markdown += f"# Страница {page['page_num']}\n\n{page['markdown']}\n\n---\n\n"
                
            mem = BytesIO()
            mem.write(full_markdown.encode('utf-8'))
            mem.seek(0)
            
            return send_file(
                mem,
                mimetype='text/markdown',
                as_attachment=True,
                download_name=f"ocr_result_{task_id}.md"
            )
    except Exception as e:
        logger.error(f"Error downloading result: {e}")
        flash(f"Ошибка при скачивании результата: {e}", 'error')
        return redirect(url_for('index'))

@app.route('/new_conversion')
def new_conversion():
    session.clear()
    return redirect(url_for('index'))

if __name__ == '__main__':
    host = '0.0.0.0'
    port = int(os.environ.get('PORT', 8011))
    app.run(debug=True, host=host, port=port)
