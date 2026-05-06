"""
Converters for OCR results to various output formats.
"""
import logging
import re
from typing import Dict, List, Any
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

from processors.handlers_registry import DocumentHandlerRegistry
from processors.table_extractor import default_extractor

logger = logging.getLogger(__name__)


def parse_ocr_result(pages: List[Dict]) -> Dict[str, Any]:
    """
    Parse OCR result pages and extract document structure.
    Uses the new document handler registry.
    """
    # Combine text from all pages
    full_text = ""
    rec_texts = []

    for page in pages:
        full_text += page.get('markdown', '') + "\n"
        if 'result_json' in page and page['result_json']:
            for p_data in page['result_json']:
                for block in p_data.get('blocks', []):
                    text = block.get('text', '')
                    if text:
                        rec_texts.append(text)
                    full_text += text + "\n"

    # Detect document type
    doc_type, type_confidence = DocumentHandlerRegistry.detect_document_type(full_text)

    if not doc_type:
        return {
            'document_type': None,
            'type_confidence': 0.0,
            'fields': [],
            'tables': [],
            'raw_text': full_text[:5000]
        }

    # Get handler and extract fields
    handler = DocumentHandlerRegistry.get_handler(doc_type)
    if not handler:
        return {
            'document_type': doc_type,
            'type_confidence': type_confidence,
            'fields': [],
            'tables': [],
            'raw_text': full_text[:5000]
        }

    # Extract regions
    regions = _extract_regions(full_text)
    regions['rec_texts'] = rec_texts

    # Parse fields
    fields = handler.extract_fields(full_text, regions)

    # Calculate average confidence
    if fields:
        type_confidence = sum(f['confidence'] for f in fields) / len(fields)

    # Extract tables
    tables = []
    for page in pages:
        # Check in page root (legacy)
        if 'tables' in page and page['tables']:
            tables.extend(page['tables'])
        # Check in result_data (VLLMProcessor)
        elif 'result_data' in page and isinstance(page['result_data'], dict) and 'tables' in page['result_data']:
            if page['result_data']['tables']:
                tables.extend(page['result_data']['tables'])

    if not tables:
        tables = default_extractor.extract_numerical_tables(full_text)

    return {
        'document_type': doc_type,
        'type_confidence': type_confidence,
        'fields': fields,
        'tables': tables,
        'raw_text': full_text[:5000]
    }


def _extract_regions(text: str) -> Dict[str, str]:
    """Extract document regions for field extraction."""
    lines = text.split('\n')
    table_start = len(lines)

    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        if any(
            re.search(p, line_lower)
            for p in [
                r'^\s*№?\s*(?:наименование|товар|описание|ед\.|кол-во|количество|сумма)',
                r'^\s*<t[dh]>.*?№.*?</t[dh]>',
                r'^\s*\d+\s+\d+\s+\d+',
                r'^\s*\|',
            ]
        ):
            table_start = i
            break

    # Find sections
    provider_start = len(lines)
    customer_start = len(lines)
    bank_start = len(lines)

    for i, line in enumerate(lines):
        line_l = line.lower()
        if provider_start == len(lines) and any(kw in line_l for kw in ['поставщик', 'продавец', 'исполнитель']):
            provider_start = i
        if customer_start == len(lines) and any(kw in line_l for kw in ['заказчик', 'покупатель', 'получатель']):
            customer_start = i
        if bank_start == len(lines) and 'банк' in line_l:
            bank_start = i

    sections = sorted([
        ('provider', provider_start),
        ('customer', customer_start),
        ('bank', bank_start),
        ('table', table_start)
    ], key=lambda x: x[1])

    result = {
        'header': '\n'.join(lines[:min(provider_start, customer_start, bank_start, table_start)]),
        'provider': '',
        'customer': '',
        'bank': '',
        'table': '\n'.join(lines[table_start:]),
    }

    for i in range(len(sections) - 1):
        name, start = sections[i]
        next_name, next_start = sections[i+1]
        if start < len(lines):
            result[name] = '\n'.join(lines[start:next_start])

    return result


def to_docx(pages: List[Dict]) -> BytesIO:
    """Convert OCR result to DOCX format with proper table support."""
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    for page in pages:
        doc.add_heading(f"Страница {page['page_num']}", level=1)
        
        result_data = page.get('result_data')
        
        if isinstance(result_data, dict):
            # Use structured data
            text = result_data.get('text', '')
            if text:
                for part in text.split('\n'):
                    if part.strip():
                        doc.add_paragraph(part.strip())
            
            tables = result_data.get('tables', [])
            if tables:
                for table_data in tables:
                    if not table_data or not isinstance(table_data, list):
                        continue
                    
                    rows_count = len(table_data)
                    cols_count = max(len(row) for row in table_data) if table_data else 0
                    
                    if rows_count > 0 and cols_count > 0:
                        doc.add_paragraph("") # Spacer
                        doc.add_heading("Таблица", level=3)
                        table = doc.add_table(rows=rows_count, cols=cols_count)
                        table.style = 'Table Grid'
                        
                        for r_idx, row_data in enumerate(table_data):
                            for c_idx, cell_data in enumerate(row_data):
                                if c_idx < cols_count:
                                    table.cell(r_idx, c_idx).text = str(cell_data)
        else:
            # Fallback to markdown text
            markdown_text = page.get('markdown', '')
            for line in markdown_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('|'):
                    # Skip table lines in fallback if they look like pipes
                    # because we don't have a good way to convert them here 
                    # without more complex logic
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph(line)

        doc.add_page_break()

    mem = BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem


def to_xlsx(pages: List[Dict]) -> BytesIO:
    """Convert OCR result to Excel format (raw data)."""
    rows = []

    for page in pages:
        if 'result_json' in page and page['result_json']:
            for p_data in page['result_json']:
                for block in p_data.get('blocks', []):
                    rows.append({
                        'Страница': page['page_num'],
                        'Текст': block.get('text', ''),
                        'Уверенность': block.get('confidence', 0)
                    })
        elif 'result_data' in page and isinstance(page['result_data'], dict) and 'text' in page['result_data']:
             rows.append({
                'Страница': page['page_num'],
                'Текст': page['result_data']['text'],
                'Уверенность': 1.0
            })
        else:
            rows.append({
                'Страница': page['page_num'],
                'Текст': page['markdown'],
                'Уверенность': 1.0
            })

    df = pd.DataFrame(rows)
    mem = BytesIO()

    with pd.ExcelWriter(mem, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='OCR Result')

    mem.seek(0)
    return mem


def to_markdown(pages: List[Dict]) -> BytesIO:
    """Convert OCR result to Markdown format."""
    full_markdown = ""

    for page in pages:
        full_markdown += f"# Страница {page['page_num']}\n\n{page['markdown']}\n\n---\n\n"

    mem = BytesIO()
    mem.write(full_markdown.encode('utf-8'))
    mem.seek(0)
    return mem


def to_extracted_data_xlsx(pages: List[Dict]) -> BytesIO:
    """Convert extracted document data to Excel with formatting."""
    parsed = parse_ocr_result(pages)

    doc_type = parsed.get('document_type', 'Неизвестный документ')
    type_confidence = parsed.get('type_confidence', 0)
    fields = parsed.get('fields', [])
    tables = parsed.get('tables', [])

    mem = BytesIO()

    with pd.ExcelWriter(mem, engine='openpyxl') as writer:
        # Styles
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Fields sheet
        if fields:
            fields_data = [
                {
                    'Поле': f['field'],
                    'Значение': f['value'],
                    'Уверенность': f['confidence'],
                    'Обязательное': 'Да' if f.get('required', False) else 'Нет'
                }
                for f in fields
            ]
            fields_df = pd.DataFrame(fields_data)
            fields_df.to_excel(writer, index=False, sheet_name='Данные документа')

            ws1 = writer.sheets['Данные документа']

            # Title
            ws1.cell(1, 1, f'Тип документа: {doc_type} (уверенность: {type_confidence:.0%})')
            ws1.merge_cells('A1:D1')
            ws1.cell(1, 1).font = Font(bold=True, size=12)

            # Header styling
            for cell in ws1[3]:
                cell.font = header_font
                cell.fill = header_fill

            # Data styling
            for row in ws1.iter_rows(min_row=3, max_row=ws1.max_row, min_col=1, max_col=4):
                for cell in row:
                    cell.border = thin_border
                    cell.alignment = Alignment(horizontal='left', vertical='center')

        # Tables sheet
        if tables:
            for idx, table_data in enumerate(tables[:10], 1):
                try:
                    table_df = pd.DataFrame(table_data)
                    sheet_name = f'Таблица {idx}'
                    table_df.to_excel(writer, index=False, sheet_name=sheet_name)
                    ws = writer.sheets[sheet_name]

                    for row in ws.iter_rows(
                        min_row=1, max_row=ws.max_row,
                        min_col=1, max_col=ws.max_column
                    ):
                        for cell in row:
                            cell.border = thin_border
                            cell.alignment = Alignment(horizontal='center', vertical='center')

                    if ws.max_row > 0:
                        for cell in ws[1]:
                            cell.font = header_font
                            cell.fill = header_fill
                except Exception as e:
                    logger.warning(f"Failed to export table {idx} to Excel: {e}")
        else:
            pd.DataFrame({'Сообщение': ['Таблицы в документе не найдены']}).to_excel(
                writer, index=False, sheet_name='Таблица 1'
            )

    mem.seek(0)
    return mem


def get_extracted_data(pages: List[Dict]) -> Dict[str, Any]:
    """Get extracted document data as JSON."""
    return parse_ocr_result(pages)
